from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "crisis-radar-guide.md"
OUTPUT = ROOT / "docs" / "Crisis_Radar_Руководство_и_методика.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = RGBColor(255, 255, 255)
BULLET_NUM_ID = 90
DECIMAL_NUM_ID = 91


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: RGBColor | None = None,
                 name: str = "Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, value, end))


def add_inline(paragraph, text: str, *, base_size: float = 11,
               base_color: RGBColor | None = None) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    for part in filter(None, pattern.split(text)):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=max(9, base_size - 1), color=DARK_BLUE, name="Courier New")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, color=base_color)


def add_numbering_definition(doc: Document, *, num_id: int, abstract_id: int,
                             number_format: str, level_text: str) -> None:
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), number_format)
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    level.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend((level, number))
    p_pr.append(num_pr)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("TRADING ASSISTANT  |  CRISIS RADAR")
    set_run_font(run, size=9, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("РУКОВОДСТВО ВЛАДЕЛЬЦА И ТЕХНИЧЕСКАЯ МЕТОДИКА")
    set_run_font(run, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Crisis Radar")
    set_run_font(run, size=30, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    run = subtitle.add_run("Как читать сигналы, как выполняются расчёты\nи где заканчиваются возможности системы")
    set_run_font(run, size=15, color=DARK_BLUE)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    callout.style = "Table Grid"
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "Рабочий сигнал: candidate-v10  •  Shadow-методика: candidate-v11\n"
        "Вероятность кризиса не публикуется, пока историческая калибровка не пройдёт защитные ворота."
    )
    set_run_font(run, size=10.5, bold=True, color=INK)

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(50)
    metadata.paragraph_format.space_after = Pt(0)
    run = metadata.add_run("Версия методики: 11 августа 2026 года")
    set_run_font(run, size=10.5, color=MUTED)

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    if columns == 8:
        widths = [2100, 900, 900, 900, 900, 900, 900, 1860]
    elif columns == 4:
        widths = [2800, 1200, 1200, 4160]
    elif columns == 7:
        widths = [2300, 900, 900, 900, 900, 900, 2560]
    elif columns == 2:
        widths = [2700, 6660]
    else:
        base = 9360 // columns
        widths = [base] * columns
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)

    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            if col_index > 0 and len(value) < 18:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(
                paragraph,
                value,
                base_size=9.5,
                base_color=INK if row_index == 0 else None,
            )
            for run in paragraph.runs:
                if row_index == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, size=8.7, color=INK, name="Courier New")
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def build() -> None:
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    add_numbering_definition(
        doc, num_id=BULLET_NUM_ID, abstract_id=90, number_format="bullet", level_text="•"
    )
    add_numbering_definition(
        doc, num_id=DECIMAL_NUM_ID, abstract_id=91, number_format="decimal", level_text="%1."
    )
    add_cover(doc)

    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(paragraph, " ".join(part.strip() for part in paragraph_buffer))
        paragraph_buffer.clear()

    in_code = False
    code_lines: list[str] = []
    skipped_title = False
    while index < len(source_lines):
        raw = source_lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines.clear()
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            rows, index = parse_table(source_lines, index)
            add_table(doc, rows)
            continue
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            title = heading.group(2)
            if level == 1 and not skipped_title:
                skipped_title = True
            else:
                doc.add_heading(title, level=min(level, 3))
            index += 1
            continue
        if re.match(r"^-\s+", stripped):
            flush_paragraph()
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, BULLET_NUM_ID)
            content = re.sub(r"^-\s+", "", stripped)
            index += 1
            while index < len(source_lines) and source_lines[index].startswith(("  ", "\t")):
                content += " " + source_lines[index].strip()
                index += 1
            add_inline(paragraph, content)
            continue
        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, DECIMAL_NUM_ID)
            content = re.sub(r"^\d+\.\s+", "", stripped)
            index += 1
            while index < len(source_lines) and source_lines[index].startswith(("  ", "\t")):
                content += " " + source_lines[index].strip()
                index += 1
            add_inline(paragraph, content)
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()

    core = doc.core_properties
    core.title = "Crisis Radar — руководство владельца и техническая методика"
    core.subject = "Исполняемая методика candidate-v10 и candidate-v11"
    core.author = "Trading Assistant"
    core.keywords = "Crisis Radar, risk monitoring, methodology, self-hosted"
    core.comments = "Сгенерировано из docs/crisis-radar-guide.md"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
