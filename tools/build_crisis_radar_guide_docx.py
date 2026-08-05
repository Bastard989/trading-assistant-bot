from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 101, 116)
INK = RGBColor(25, 34, 47)
TABLE_FILL = "E8EEF5"
WIDTH_DXA = 9360


def font(run, *, size: float, color=INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_header_row(row) -> None:
    row_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_pr.append(header)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            font(run, size=10, color=DARK_BLUE)
            run.font.name = "Menlo"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Menlo")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            font(run, size=11, bold=True)
        else:
            run = paragraph.add_run(part)
            font(run, size=11)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shading)
    run = paragraph.add_run("\n".join(lines))
    font(run, size=9.5, color=DARK_BLUE)
    run.font.name = "Menlo"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Menlo")


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [rows[0], *rows[2:]]
    columns = len(rows[0])
    widths = [2500, 1000, 1000, 1000, 3860] if columns == 5 else [WIDTH_DXA // columns] * columns
    widths[-1] += WIDTH_DXA - sum(widths)
    body_rows = rows[1:]
    body_chunks = [body_rows] if len(body_rows) <= 10 else [body_rows[:8], body_rows[8:]]
    for chunk_index, body in enumerate(body_chunks):
        if chunk_index:
            document.add_page_break()
        table_rows = [rows[0], *body]
        table = document.add_table(rows=len(table_rows), cols=columns)
        table.style = "Table Grid"
        for row_index, values in enumerate(table_rows):
            for column_index, value in enumerate(values):
                cell = table.cell(row_index, column_index)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run(value)
                font(run, size=9 if columns >= 5 else 10, bold=row_index == 0)
                if row_index == 0:
                    set_cell_shading(cell, TABLE_FILL)
        set_table_geometry(table, widths)
        mark_header_row(table.rows[0])
        document.add_paragraph().paragraph_format.space_after = Pt(2)


def build(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.text = "TRADING ASSISTANT  /  CRISIS RADAR"
    font(header.runs[0], size=8.5, color=MUTED, bold=True)
    add_page_field(section.footer.paragraphs[0])

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(110)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Crisis Radar")
    font(run, size=30, color=DARK_BLUE, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Руководство пользователя и техническая методика")
    font(run, size=15, color=BLUE)
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(110)
    run = metadata.add_run("candidate-v10  •  scenario-fusion-v1  •  4 августа 2026")
    font(run, size=10, color=MUTED, italic=True)
    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lead.add_run("Проверяемая система раннего предупреждения — не гарантия кризиса и не команда на сделку")
    font(run, size=11, color=INK, bold=True)
    document.add_section(WD_SECTION.NEW_PAGE)

    index = 1  # source title is represented by the cover
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("| ") and index + 1 < len(lines) and lines[index + 1].startswith("|---"):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue
        if line.startswith("## "):
            document.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            document.add_paragraph(line[4:], style="Heading 2")
        elif re.match(r"^\d+\. ", line):
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, "\u00a0" + re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, "\u00a0" + line[2:])
        elif line.strip():
            paragraph = document.add_paragraph()
            add_inline(paragraph, line)
        index += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_crisis_radar_guide_docx.py SOURCE.md OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
